from __future__ import annotations

from pathlib import Path

import pytest

from byzantine.citations import format_chicago_note, format_gbt7714
from byzantine.models.document import BibliographicMetadata
from byzantine.retrieval.hybrid import reciprocal_rank_fusion
from byzantine.storage.database import LibraryDatabase
from byzantine.workflows.delete_document import delete_document_from_library
from byzantine.workflows.process_document import (
    _extract_docx,
    _extract_pdf,
    process_document,
    reprocess_document,
)


def _document(database: LibraryDatabase, file_hash: str = "a" * 64):
    return database.create_document(
        collection_id="personal",
        metadata=BibliographicMetadata(
            title="Anna Komnene",
            author="Anna Komnene",
            edition="2nd ed.",
            language="English",
            source_type="translation",
        ),
        file_path="C:/library/alexiad.txt",
        file_hash=file_hash,
        mime_type="text/plain",
    )


def _chunk(document_id: str) -> dict[str, object]:
    return {
        "chunk_id": f"{document_id}_chunk_00000",
        "chunk_index": 0,
        "section_path": ["Book I", "Alexios"],
        "text": "Alexios prepared the army at Constantinople.",
        "search_text": "Alexios prepared the army at Constantinople.",
        "page_start": 12,
        "page_end": 12,
        "source_regions": [
            {
                "page": 12,
                "region_id": "region_12_03",
                "bbox": [82, 315, 506, 468],
                "coordinate_space": "pdf_points",
                "page_width": 595,
                "page_height": 842,
            }
        ],
        "metadata": {"people": ["Alexios I"], "places": ["Constantinople"], "topics": ["warfare"]},
    }


def test_idempotent_library_and_fts_evidence(tmp_path):
    database = LibraryDatabase(tmp_path / "library.db")
    database.initialize()
    database.initialize()
    assert {item["collection_id"] for item in database.collections()} == {"starter", "personal"}

    document = _document(database)
    database.save_chunks(document.document_id, [_chunk(document.document_id)])
    rows = database.fts_search("Alexios")
    assert len(rows) == 1
    evidence = database.evidence_from_row(rows[0])
    assert evidence.document_id == document.document_id
    assert evidence.source_regions[0].bbox == [82.0, 315.0, 506.0, 468.0]
    assert "Anna Komnene" in format_gbt7714(evidence)
    assert "Anna Komnene" in format_chicago_note(evidence)


def test_documents_do_not_overwrite_and_duplicate_is_rejected(tmp_path):
    database = LibraryDatabase(tmp_path / "library.db")
    database.initialize()
    first = _document(database, "1" * 64)
    second = _document(database, "2" * 64)
    database.save_chunks(first.document_id, [_chunk(first.document_id)])
    database.save_chunks(second.document_id, [_chunk(second.document_id)])
    assert len(database.document_evidence(first.document_id)) == 1
    assert len(database.document_evidence(second.document_id)) == 1
    with pytest.raises(ValueError, match="已经导入"):
        _document(database, "1" * 64)


def test_contradiction_history_returns_parsed_evidence(tmp_path):
    database = LibraryDatabase(tmp_path / "library.db")
    database.initialize()
    document = _document(database)
    database.save_chunks(document.document_id, [_chunk(document.document_id)])
    evidence = database.document_evidence(document.document_id)[0]

    database.save_contradiction(
        subject="Alexios prepared the army.",
        description="A claim and its counter-reading.",
        classification="perspective",
        evidence_side_a=evidence,
        evidence_side_b=evidence,
    )

    history = database.list_contradictions()
    assert history[0]["subject"] == "Alexios prepared the army."
    assert history[0]["evidence_side_a"]["document_id"] == document.document_id


def test_rrf_deduplicates_and_favors_consistent_results():
    assert reciprocal_rank_fusion([["a", "b"], ["b", "c"]])[:3] == ["b", "a", "c"]


def test_txt_import_preserves_document_scope_and_text_regions(tmp_path, monkeypatch):
    from byzantine.indexing import library_index

    monkeypatch.setenv("BYZANTINE_DATA_DIR", str(tmp_path / "app-data"))
    monkeypatch.setattr(library_index, "upsert_evidence", lambda *args, **kwargs: None)
    source = tmp_path / "chronicle.txt"
    source.write_text(
        "Alexios arrived at Constantinople.\n\nThe army was prepared.", encoding="utf-8"
    )
    database = LibraryDatabase(tmp_path / "app-data" / "library.db")
    document = process_document(
        source,
        collection_id="personal",
        metadata=BibliographicMetadata(title="Chronicle", language="English"),
        database=database,
    )
    evidence = database.document_evidence(document.document_id)
    assert document.status == "ready"
    assert evidence[0].document_id == document.document_id
    assert evidence[0].source_regions[0].coordinate_space == "text_characters"


def test_pdf_import_persists_page_and_bbox(tmp_path, monkeypatch):
    import fitz

    from byzantine.indexing import library_index

    monkeypatch.setenv("BYZANTINE_DATA_DIR", str(tmp_path / "app-data"))
    monkeypatch.setattr(library_index, "upsert_evidence", lambda *args, **kwargs: None)
    source = tmp_path / "source.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Alexios governed Constantinople.")
    pdf.save(source)
    pdf.close()
    database = LibraryDatabase(tmp_path / "app-data" / "library.db")
    document = process_document(
        source,
        collection_id="starter",
        metadata=BibliographicMetadata(title="Test PDF", language="English"),
        database=database,
    )
    region = database.document_evidence(document.document_id)[0].source_regions[0]
    assert region.page == 1
    assert region.coordinate_space == "pdf_points"
    assert region.bbox is not None


def test_pdf_layout_blocks_are_aggregated_into_one_evidence_chunk(tmp_path):
    import fitz

    source = tmp_path / "blocks.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    for y, text in (
        (72, "First paragraph."),
        (120, "Second paragraph."),
        (168, "Third paragraph."),
    ):
        page.insert_text((72, y), text)
    pdf.save(source)
    pdf.close()
    chunks, page_count = _extract_pdf(source, "doc_test")
    assert page_count == 1
    assert len(chunks) == 1
    assert len(chunks[0]["source_regions"]) == 3


def test_docx_import_preserves_headings_and_tables(tmp_path):
    from docx import Document as WordDocument

    source = tmp_path / "chronicle.docx"
    word = WordDocument()
    word.add_heading("Book I", level=1)
    word.add_paragraph("Alexios prepared the army at Constantinople.")
    table = word.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Year"
    table.cell(0, 1).text = "Event"
    word.save(source)

    chunks, page_count = _extract_docx(source, "doc_test")

    assert page_count is None
    assert chunks[0]["section_path"] == ["Book I"]
    assert chunks[0]["source_regions"][0]["coordinate_space"] == "docx_paragraphs"
    assert "Year | Event" in chunks[-1]["text"]


def test_reprocess_recovers_legacy_missing_file_path(tmp_path, monkeypatch):
    from byzantine.indexing import library_index

    monkeypatch.setenv("BYZANTINE_DATA_DIR", str(tmp_path / "app-data"))
    monkeypatch.setattr(library_index, "upsert_evidence", lambda *args, **kwargs: None)
    monkeypatch.setattr(library_index, "delete_evidence", lambda *args, **kwargs: None)
    source = tmp_path / "chronicle.txt"
    source.write_text("A source paragraph about Constantinople.", encoding="utf-8")
    database = LibraryDatabase(tmp_path / "app-data" / "library.db")
    original = process_document(
        source,
        collection_id="personal",
        metadata=BibliographicMetadata(title="Chronicle"),
        database=database,
    )
    database.update_document(original.document_id, file_path="")

    repaired = reprocess_document(original.document_id, database=database)
    assert repaired.file_path.endswith("source.txt")
    assert Path(repaired.file_path).is_file()

    database.update_document(original.document_id, file_path="")
    duplicate_retry = process_document(
        source,
        collection_id="personal",
        metadata=BibliographicMetadata(title="Chronicle"),
        database=database,
    )
    assert duplicate_retry.document_id == original.document_id
    assert Path(duplicate_retry.file_path).is_file()


def test_delete_document_removes_vectors_records_and_stored_copy(tmp_path, monkeypatch):
    from byzantine.indexing import library_index

    root = tmp_path / "app-data"
    monkeypatch.setenv("BYZANTINE_DATA_DIR", str(root))
    deleted_vectors: list[str] = []
    monkeypatch.setattr(
        library_index,
        "delete_evidence",
        lambda chunk_ids, **kwargs: deleted_vectors.extend(chunk_ids),
    )
    database = LibraryDatabase(root / "library.db")
    database.initialize()
    document = _document(database)
    stored = root / "documents" / document.document_id
    stored.mkdir(parents=True)
    (stored / "source.txt").write_text("source", encoding="utf-8")
    database.update_document(document.document_id, file_path=str(stored / "source.txt"))
    database.save_chunks(document.document_id, [_chunk(document.document_id)])
    evidence = database.document_evidence(document.document_id)[0]
    topic = database.create_topic("Test topic")
    database.add_topic_item(topic, "answer", evidence=[evidence])
    conversation = database.create_conversation(
        title="Alexios chat",
        collection_ids=["personal"],
        document_ids=[document.document_id],
        topic_id=topic,
    )
    database.add_chat_message(conversation, role="user", content="What did Alexios do?")
    database.add_chat_message(
        conversation, role="assistant", content="He prepared the army.", evidence=[evidence]
    )
    database.save_topic_chat_summary(
        topic_id=topic,
        conversation_id=conversation,
        title="Military preparation",
        tags=["Alexios", "warfare"],
        summary="Alexios prepared the army at Constantinople [E1].",
        evidence=[evidence],
    )
    claim = database.create_claim("Test claim")
    database.link_claim_evidence(claim, evidence, "support")

    delete_document_from_library(document.document_id, database=database)

    assert deleted_vectors == [evidence.chunk_id]
    assert not stored.exists()
    assert database.list_documents() == []
    assert database.fts_search("Alexios") == []
    with database.connect() as connection:
        assert connection.execute("SELECT count(*) FROM topic_items").fetchone()[0] == 0
        assert connection.execute("SELECT count(*) FROM claim_evidence").fetchone()[0] == 0
        assert connection.execute("SELECT count(*) FROM conversations").fetchone()[0] == 0
        assert connection.execute("SELECT count(*) FROM chat_messages").fetchone()[0] == 0
        assert connection.execute("SELECT count(*) FROM topic_chat_summaries").fetchone()[0] == 0


def test_topic_chat_summary_preserves_conversation_and_evidence(tmp_path):
    database = LibraryDatabase(tmp_path / "library.db")
    database.initialize()
    document = _document(database)
    database.save_chunks(document.document_id, [_chunk(document.document_id)])
    evidence = database.document_evidence(document.document_id)[0]
    topic = database.create_topic("Military districts")
    conversation = database.create_conversation(
        title="Themes question",
        collection_ids=["personal"],
        document_ids=[document.document_id],
    )
    database.add_chat_message(conversation, role="user", content="How did the themes change?")
    database.add_chat_message(
        conversation,
        role="assistant",
        content="The evidence describes military preparation.",
        evidence=[evidence],
        labels=["warfare"],
    )
    database.save_topic_chat_summary(
        topic_id=topic,
        conversation_id=conversation,
        title="Evolution of the themes",
        tags=["themes", "warfare"],
        summary="The conversation discusses the military system [E1].",
        evidence=[evidence],
    )

    saved = database.topic_chat_summaries(topic)
    assert saved[0]["title"] == "Evolution of the themes"
    assert saved[0]["tags"] == ["themes", "warfare"]
    assert saved[0]["evidence_snapshot"][0]["document_id"] == document.document_id
    assert database.get_conversation(conversation)["topic_id"] == topic
    assert len(database.conversation_messages(conversation)) == 2
