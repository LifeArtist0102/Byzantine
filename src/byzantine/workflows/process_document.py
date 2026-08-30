"""Local import pipeline with recoverable stages and progress callbacks."""

from __future__ import annotations

import hashlib
import mimetypes
import os
import re
import shutil
from collections.abc import Callable
from pathlib import Path
from typing import Any

from byzantine.chunking.semantic import build_hierarchical_chunks
from byzantine.metadata.enrichment import build_retrieval_text, enrich_chunk, load_seed
from byzantine.metadata.inference import infer_metadata_batches
from byzantine.models.document import BibliographicMetadata, DocumentRecord
from byzantine.paths import ensure_app_data_dir
from byzantine.storage.database import LibraryDatabase

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".markdown", ".jpg", ".jpeg", ".png"}
ProgressCallback = Callable[[str, float], None]


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _link_chunks(chunks: list[dict[str, Any]], document_id: str) -> list[dict[str, Any]]:
    for index, chunk in enumerate(chunks):
        chunk["chunk_id"] = f"{document_id}_chunk_{index:05d}"
        chunk["chunk_index"] = index
        chunk["prev_chunk_id"] = chunks[index - 1]["chunk_id"] if index else None
        chunk["next_chunk_id"] = chunks[index + 1]["chunk_id"] if index + 1 < len(chunks) else None
    return chunks


def _local_semantic_embedder() -> Callable[[list[str]], list[list[float]]] | None:
    """Use a local BGE-M3 model for boundary hints, never an online model."""
    if os.getenv("BYZANTINE_SEMANTIC_BOUNDARIES", "1") == "0":
        return None
    model_path = os.getenv("BYZANTINE_EMBEDDING_MODEL")
    bundled = Path(__file__).resolve().parents[3] / "models" / "bge-m3"
    if not model_path and bundled.is_dir():
        model_path = str(bundled)
    if not model_path:
        return None
    try:
        import torch
        from FlagEmbedding import BGEM3FlagModel
    except ImportError:
        return None
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = BGEM3FlagModel(model_path, use_fp16=device == "cuda", device=device)

    def embed(texts: list[str]) -> list[list[float]]:
        return model.encode(
            texts,
            batch_size=8,
            max_length=1024,
            return_dense=True,
            return_sparse=False,
            return_colbert_vecs=False,
        )["dense_vecs"].tolist()

    return embed


def _text_chunks(
    text: str,
    *,
    document_id: str,
    source_regions: list[dict[str, Any]],
    page: int | None = None,
) -> list[dict[str, Any]]:
    paragraphs = [part.strip() for part in text.replace("\r\n", "\n").split("\n\n") if part.strip()]
    chunks: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        for start in range(0, len(paragraph), 2600):
            part = paragraph[start : start + 2600]
            chunks.append(
                {
                    "chunk_id": "",
                    "chunk_index": 0,
                    "section_path": ["Imported document"],
                    "text": part,
                    "search_text": part,
                    "page_start": page,
                    "page_end": page,
                    "source_regions": source_regions,
                    "metadata": {},
                }
            )
    return _link_chunks(chunks, document_id)


def _extract_pdf(path: Path, document_id: str) -> tuple[list[dict[str, Any]], int]:
    try:
        import fitz
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('请安装 PyMuPDF：pip install -e ".[app]"') from exc
    chunks: list[dict[str, Any]] = []
    section_path = ["PDF document"]
    numbered_heading = re.compile(r"^(?:\d+(?:\.\d+)*|[IVXLC]+)\.?\s+", re.IGNORECASE)
    with fitz.open(path) as pdf:
        for page_number, page in enumerate(pdf, start=1):
            page_dict = page.get_text("dict")
            text_blocks = sorted(
                [block for block in page_dict.get("blocks", []) if block.get("type") == 0],
                key=lambda block: (block["bbox"][1], block["bbox"][0]),
            )
            span_sizes = [
                span.get("size", 0.0)
                for block in text_blocks
                for line in block.get("lines", [])
                for span in line.get("spans", [])
                if span.get("text", "").strip()
            ]
            body_size = sorted(span_sizes)[len(span_sizes) // 2] if span_sizes else 0.0
            buffer: list[str] = []
            regions: list[dict[str, Any]] = []

            def emit(
                path: list[str],
                buffer: list[str] = buffer,
                regions: list[dict[str, Any]] = regions,
                page_number: int = page_number,
            ) -> None:
                if not buffer:
                    return
                text = "\n\n".join(buffer).strip()
                if text:
                    chunks.append(
                        {
                            "chunk_id": "",
                            "chunk_index": 0,
                            "section_path": path.copy(),
                            "text": text,
                            "search_text": text,
                            "page_start": page_number,
                            "page_end": page_number,
                            "source_regions": regions.copy(),
                            "metadata": {},
                        }
                    )
                buffer.clear()
                regions.clear()

            for block_index, block in enumerate(text_blocks):
                x0, y0, x1, y1 = block["bbox"]
                spans = [
                    span
                    for line in block.get("lines", [])
                    for span in line.get("spans", [])
                    if span.get("text", "").strip()
                ]
                text = " ".join(str(span["text"]).strip() for span in spans).strip()
                if not text:
                    continue
                max_size = max((float(span.get("size", 0.0)) for span in spans), default=0.0)
                bold = any("bold" in str(span.get("font", "")).casefold() for span in spans)
                word_count = len(text.split())
                is_heading = word_count <= 28 and (
                    bool(numbered_heading.match(text))
                    or (body_size and max_size >= body_size * 1.18)
                    or (bold and max_size >= body_size)
                )
                if is_heading:
                    emit(section_path)
                    # Numbered headings carry a conservative hierarchy; other
                    # typographic headings begin a new section at this level.
                    number = numbered_heading.match(text)
                    level = number.group(0).count(".") + 1 if number else 1
                    section_path = [*section_path[:level], text]
                    continue
                region = {
                    "page": page_number,
                    "region_id": f"region_{page_number}_{block_index}",
                    "bbox": [x0, y0, x1, y1],
                    "coordinate_space": "pdf_points",
                    "page_width": page.rect.width,
                    "page_height": page.rect.height,
                }
                if buffer and len("\n\n".join(buffer)) + len(text) + 2 > 2200:
                    emit(section_path)
                for start in range(0, len(text), 2600):
                    part = text[start : start + 2600]
                    if buffer and len("\n\n".join(buffer)) + len(part) + 2 > 2200:
                        emit(section_path)
                    buffer.append(part)
                    regions.append(region)
                    if len(part) >= 2600:
                        emit(section_path)
            emit(section_path)
        return _link_chunks(chunks, document_id), len(pdf)


def _extract_text(path: Path, document_id: str) -> tuple[list[dict[str, Any]], int | None]:
    text = path.read_text(encoding="utf-8", errors="replace")
    regions = [
        {
            "region_id": "paragraph_source",
            "coordinate_space": "text_characters",
            "paragraph_index": 0,
            "character_start": 0,
            "character_end": len(text),
        }
    ]
    return _text_chunks(text, document_id=document_id, source_regions=regions), None


def _extract_docx(path: Path, document_id: str) -> tuple[list[dict[str, Any]], int | None]:
    """Extract paragraphs and tables while preserving Word heading structure."""
    try:
        from docx import Document as WordDocument
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("DOCX 导入需要 python-docx，请重新安装项目依赖。") from exc

    document = WordDocument(path)
    chunks: list[dict[str, Any]] = []
    section_path = ["Imported DOCX"]

    def append_text(text: str, *, paragraph_index: int, path_parts: list[str]) -> None:
        clean = text.strip()
        if not clean:
            return
        for start in range(0, len(clean), 2600):
            part = clean[start : start + 2600]
            chunks.append(
                {
                    "chunk_id": "",
                    "chunk_index": 0,
                    "section_path": path_parts.copy(),
                    "text": part,
                    "search_text": part,
                    "page_start": None,
                    "page_end": None,
                    "source_regions": [
                        {
                            "region_id": f"docx_paragraph_{paragraph_index}_{start}",
                            "coordinate_space": "docx_paragraphs",
                            "paragraph_index": paragraph_index,
                            "character_start": start,
                            "character_end": start + len(part),
                        }
                    ],
                    "metadata": {},
                }
            )

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if text and style_name.lower().startswith("heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1
            section_path = section_path[: max(level - 1, 0)]
            section_path.append(text)
            continue
        append_text(text, paragraph_index=paragraph_index, path_parts=section_path)

    paragraph_offset = len(document.paragraphs)
    for table_index, table in enumerate(document.tables):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        append_text(
            "\n".join(row for row in rows if row.strip(" |")),
            paragraph_index=paragraph_offset + table_index,
            path_parts=[*section_path, f"Table {table_index + 1}"],
        )
    return _link_chunks(chunks, document_id), None


def _extract_image(path: Path, document_id: str) -> tuple[list[dict[str, Any]], int | None]:
    try:
        from paddleocr import PaddleOCR
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError('图片或扫描件需要 OCR。请安装：pip install -e ".[ocr]"') from exc
    result = PaddleOCR(use_angle_cls=True, lang="en").ocr(str(path), cls=True)
    chunks: list[dict[str, Any]] = []
    for item in result[0] if result else []:
        polygon, (text, confidence) = item
        xs, ys = [point[0] for point in polygon], [point[1] for point in polygon]
        region = {
            "region_id": f"ocr_{len(chunks):04d}",
            "coordinate_space": "image_pixels",
            "bbox": [min(xs), min(ys), max(xs), max(ys)],
        }
        chunks.extend(_text_chunks(text, document_id=document_id, source_regions=[region]))
        chunks[-1]["metadata"] = {
            "ocr_raw": text,
            "corrected_text": text,
            "ocr_confidence": confidence,
        }
    return _link_chunks(chunks, document_id), 1


def _run_pipeline(
    *,
    source: Path,
    document: DocumentRecord,
    database: LibraryDatabase,
    root: Path,
    seed_path: Path | None,
    progress: ProgressCallback | None,
) -> DocumentRecord:
    database.update_document(document.document_id, status="extracting", error_message=None)
    if progress:
        progress("正在提取文字与原文坐标", 0.08)
    if source.suffix.lower() == ".pdf":
        chunks, page_count = _extract_pdf(source, document.document_id)
    elif source.suffix.lower() == ".docx":
        chunks, page_count = _extract_docx(source, document.document_id)
    elif source.suffix.lower() in {".txt", ".md", ".markdown"}:
        chunks, page_count = _extract_text(source, document.document_id)
    else:
        database.update_document(document.document_id, status="ocr_processing")
        if progress:
            progress("正在进行 OCR 识别", 0.15)
        chunks, page_count = _extract_image(source, document.document_id)
    if not chunks:
        raise ValueError("未从文件中提取到可检索文本。")
    database.update_document(document.document_id, status="enriching", page_count=page_count)
    if progress:
        progress("正在聚合段落并生成检索元数据", 0.30)
    seed = load_seed(seed_path) if seed_path else {}
    chunks = build_hierarchical_chunks(
        chunks,
        document_id=document.document_id,
        semantic_embedder=_local_semantic_embedder(),
    )
    # Deterministic enrichment always happens first.  Optional LLM metadata
    # only augments the difficult children selected by its gate below.
    enriched = [enrich_chunk(chunk, seed) for chunk in chunks]
    bibliographic = document.model_dump(mode="json")
    trusted_bibliographic = {
        key: bibliographic.get(key)
        for key in (
            "title",
            "author",
            "translator",
            "edition",
            "publisher",
            "publication_year",
            "language",
            "source_type",
        )
        if bibliographic.get(key) not in (None, "")
    }
    for item in enriched:
        item["metadata"].setdefault("trusted", {})["bibliographic"] = trusted_bibliographic
    if os.getenv("BYZANTINE_ENABLE_METADATA_LLM", "0") == "1":
        api_key = os.getenv("DEEPSEEK_API_KEY")
        if api_key:
            try:
                from openai import OpenAI

                enriched = infer_metadata_batches(
                    enriched,
                    database=database,
                    document_hash=document.file_hash,
                    model_name=os.getenv("BYZANTINE_METADATA_MODEL", "deepseek-chat"),
                    client=OpenAI(api_key=api_key, base_url="https://api.deepseek.com"),
                )
            except (ImportError, RuntimeError, ValueError):
                # Importing a document must not fail merely because optional
                # retrieval metadata could not be inferred.
                pass
    finalized: list[dict[str, Any]] = []
    for item in enriched:
        metadata = item["metadata"]
        item["retrieval_text"] = build_retrieval_text(
            item["original_text"],
            bibliographic=bibliographic,
            section_path=item["section_path"],
            trusted=dict(metadata.get("trusted", {})),
            candidates=dict(metadata.get("candidate", {})),
            llm=dict(metadata.get("llm_inference", {})),
        )
        item["search_text"] = item["retrieval_text"]
        finalized.append(item)
    chunks = finalized
    try:
        from byzantine.indexing.library_index import delete_evidence

        delete_evidence(
            [item.chunk_id for item in database.document_evidence(document.document_id)],
            qdrant_path=str(root / "qdrant"),
        )
    except RuntimeError:
        pass
    database.save_chunks(document.document_id, chunks)
    try:
        from byzantine.indexing.library_index import upsert_evidence

        evidence = database.document_evidence(document.document_id)
        database.update_document(document.document_id, status="indexing")
        if progress:
            progress(f"正在向量化 {len(evidence)} 条证据", 0.42)
        upsert_evidence(
            evidence,
            qdrant_path=str(root / "qdrant"),
            progress=(
                lambda completed, total: (
                    progress(
                        f"正在向量化 {completed}/{total} 条证据",
                        0.42 + 0.55 * completed / max(total, 1),
                    )
                    if progress
                    else None
                )
            ),
        )
    except RuntimeError as exc:
        database.update_document(document.document_id, error_message=f"向量索引未建立：{exc}")
    database.update_document(document.document_id, status="ready", page_count=page_count)
    if progress:
        progress("处理完成", 1.0)
    return database.get_document(document.document_id)


def process_document(
    source: Path,
    *,
    collection_id: str,
    metadata: BibliographicMetadata,
    database: LibraryDatabase | None = None,
    seed_path: Path | None = None,
    progress: ProgressCallback | None = None,
) -> DocumentRecord:
    """Synchronously import one document with durable source provenance."""
    if source.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise ValueError("仅支持 PDF、DOCX、TXT、Markdown、JPG、JPEG、PNG。")
    root = ensure_app_data_dir()
    database = database or LibraryDatabase(root / "library.db")
    database.initialize()
    digest = file_hash(source)
    duplicate = database.find_duplicate(digest)
    if duplicate:
        document = database.get_document(str(duplicate["document_id"]))
        if document.status == "ready" and Path(document.file_path).is_file():
            raise ValueError(f"该文件已经导入：{document.title}（{document.document_id}）")
        destination = root / "documents" / document.document_id
        destination.mkdir(parents=True, exist_ok=True)
        stored = destination / f"source{source.suffix.lower()}"
        shutil.copy2(source, stored)
        database.update_document(document.document_id, file_path=str(stored))
        return reprocess_document(
            document.document_id, database=database, seed_path=seed_path, progress=progress
        )
    document = database.create_document(
        collection_id=collection_id,
        metadata=metadata,
        file_path="",
        file_hash=digest,
        mime_type=mimetypes.guess_type(source.name)[0] or "application/octet-stream",
    )
    destination = root / "documents" / document.document_id
    destination.mkdir(parents=True, exist_ok=True)
    stored = destination / f"source{source.suffix.lower()}"
    shutil.copy2(source, stored)
    database.update_document(document.document_id, file_path=str(stored))
    try:
        return _run_pipeline(
            source=stored,
            document=document,
            database=database,
            root=root,
            seed_path=seed_path,
            progress=progress,
        )
    except Exception as exc:
        database.update_document(document.document_id, status="failed", error_message=str(exc))
        raise


def reprocess_document(
    document_id: str,
    *,
    database: LibraryDatabase,
    seed_path: Path | None = None,
    progress: ProgressCallback | None = None,
) -> DocumentRecord:
    """Retry one source without changing its document ID."""
    document = database.get_document(document_id)
    root = ensure_app_data_dir()
    source = Path(document.file_path)
    if not source.is_file():
        candidates = sorted((root / "documents" / document_id).glob("source.*"))
        if not candidates:
            raise FileNotFoundError("找不到本地原文件，无法重新处理。请重新上传该文献。")
        source = candidates[0]
        database.update_document(document_id, file_path=str(source))
    try:
        return _run_pipeline(
            source=source,
            document=document,
            database=database,
            root=root,
            seed_path=seed_path,
            progress=progress,
        )
    except Exception as exc:
        database.update_document(document_id, status="failed", error_message=str(exc))
        raise
